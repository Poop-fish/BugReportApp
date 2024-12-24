import sys
import sqlite3
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QTextEdit, QPushButton, QComboBox, QTableWidget,
    QTableWidgetItem, QDateEdit, QSplitter, QStatusBar, QFileDialog
)
from PyQt5.QtCore import Qt, QDate
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class BugReportApp(QMainWindow):
    def __init__(SELF):
        super().__init__()
        SELF.setWindowTitle("Game Bug Reporter")
        SELF.setGeometry(100, 100, 900, 600)
        SELF.selected_bug_id = None
        SELF.init_ui()
        SELF.init_database()

    def init_ui(SELF):
        SELF.central_widget = QWidget()
        SELF.setCentralWidget(SELF.central_widget)
        layout = QVBoxLayout()

        SELF.setStyleSheet("""
            QWidget {
                background: linear-gradient(to bottom, #000000, #000000);  /* Black background */
                color: #fc0303;  /* White text color */
            }
        """)

        SELF.splitter = QSplitter(Qt.Vertical)
        SELF.splitter.setStyleSheet("""
            QSplitter::handle {
                background-color: #444444;  /* Dark handle */
                border: 2px solid #39FF14;  /* Neon green border */
            }
        """)

        SELF.form_widget = QWidget()
        FORM_LAYOUT = QVBoxLayout()

        # Name Input Field
        SELF.NAME_INPUT = QLineEdit()
        SELF.NAME_INPUT.setPlaceholderText("Enter Names Here .... ")
        SELF.NAME_INPUT.setStyleSheet("""
            QLineEdit {
                background-color: rgba(0, 0, 0, 0.5);  /* Dark transparent background */
                border: 2px solid #39FF14;  /* Neon green border */
                border-radius: 5px;
                color: cyan;  /* White text */
                padding: 10px;
            }
            QLineEdit:focus {
                border-color: #FF0000;  /* Red border when focused */
            }
        """)
        FORM_LAYOUT.addWidget(QLabel("Names Of Who Found The Bug:"))
        FORM_LAYOUT.addWidget(SELF.NAME_INPUT)

        # Title Input Field
        SELF.TITLE_INPUT = QLineEdit()
        SELF.TITLE_INPUT.setPlaceholderText("Enter Bug Title Text Here ...")
        SELF.TITLE_INPUT.setStyleSheet("""
            QLineEdit {
                background-color: rgba(0, 0, 0, 0.5);
                border: 2px solid #39FF14;  /* Neon green border */
                border-radius: 5px;
                color: lime;
                padding: 10px;
            }
            QLineEdit:focus {
                border-color: #FF0000;  /* Red border when focused */
            }
        """)
        FORM_LAYOUT.addWidget(QLabel("Bug Title:"))
        FORM_LAYOUT.addWidget(SELF.TITLE_INPUT)

        # Description Input Field
        SELF.DESCRIPTION_INPUT = QTextEdit()
        SELF.DESCRIPTION_INPUT.setPlaceholderText("Enter detailed description of the bug you found...")
        SELF.DESCRIPTION_INPUT.setStyleSheet("""
            QTextEdit {
                background-color: rgba(0, 0, 0, 0.5);
                border: 2px solid #39FF14;  /* Neon green border */
                border-radius: 5px;
                color: blue;
                padding: 10px;
            }
            QTextEdit:focus {
                border-color: #FF0000;  /* Red border when focused */
            }
        """)
        FORM_LAYOUT.addWidget(QLabel("Bug Description:"))
        FORM_LAYOUT.addWidget(SELF.DESCRIPTION_INPUT)

        # Severity Dropdown
        SELF.SEVERITY_DROPDOWN = QComboBox()
        SELF.SEVERITY_DROPDOWN.addItems(["Low", "Medium", "High", "Critical"])
        SELF.SEVERITY_DROPDOWN.setStyleSheet("""
            QComboBox {
                background-color: rgba(0, 0, 0, 0.5);
                border: 2px solid #39FF14;  /* Neon green border */
                border-radius: 5px;
                color: yellow;
                padding: 5px;
            }
            QComboBox:focus {
                border-color: #FF0000;  /* Red border when focused */
            }
        """)
        FORM_LAYOUT.addWidget(QLabel("Select Severity:"))
        FORM_LAYOUT.addWidget(SELF.SEVERITY_DROPDOWN)

        # Date Input Field
        SELF.DATE_INPUT = QDateEdit()
        SELF.DATE_INPUT.setCalendarPopup(True)
        SELF.DATE_INPUT.setDate(QDate.currentDate())
        SELF.DATE_INPUT.setStyleSheet("""
            QDateEdit {
                background-color: rgba(0, 0, 0, 0.5);
                border: 2px solid #39FF14;  /* Neon green border */
                border-radius: 5px;
                color: black;
                padding: 5px;
            }
            QDateEdit:focus {
                border-color: #FF0000;  /* Red border when focused */
            }
        """)
        FORM_LAYOUT.addWidget(QLabel("Select Submission Date:"))
        FORM_LAYOUT.addWidget(SELF.DATE_INPUT)

        # Button Layout
        BUTTON_LAYOUT = QHBoxLayout()

        # Submit Button
        SELF.SUBMIT_BUTTON = QPushButton("Submit Bug")
        SELF.SUBMIT_BUTTON.setStyleSheet("""
            QPushButton {
                background-color: #39FF14;  /* Neon green background */
                border: 2px solid #39FF14;  /* Neon green border */
                color: #000000;  /* Black text */
                padding: 10px;
                border-radius: 10px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #FF0000;  /* Red background on hover */
                border: 2px solid #FF0000;  /* Red border on hover */
                color: #FFFFFF;  /* White text on hover */
            }
        """)
        SELF.SUBMIT_BUTTON.clicked.connect(SELF.submit_bug)
        BUTTON_LAYOUT.addWidget(SELF.SUBMIT_BUTTON)

        # Mark as Done Button
        SELF.MARK_DONE_BUTTON = QPushButton("Mark as Done")
        SELF.MARK_DONE_BUTTON.setStyleSheet("""
    QPushButton {
        background-color: #FF0000;  /* Red background */
        border: 2px solid #FF0000;  /* Red border */
        color: #FFFFFF;  /* White text */
        padding: 10px;
        border-radius: 10px;
        font-weight: bold;
    }
    QPushButton:hover {
        background-color: blue;  /* Neon green background on hover */
        border: 2px solid blue;  /* Neon green border on hover */
        color: #000000;  /* Black text on hover */
    }
    QPushButton:disabled {
        background-color: #FFCCCC;  /* Light red background for disabled state */
        color: #999999;  /* Grey text for disabled state */
        border: 2px solid #FFCCCC;  /* Light red border for disabled state */
    }
""")
        SELF.MARK_DONE_BUTTON.clicked.connect(SELF.mark_bug_done)
        SELF.MARK_DONE_BUTTON.setEnabled(False)
        BUTTON_LAYOUT.addWidget(SELF.MARK_DONE_BUTTON)

        # Save Button
        SELF.SAVE_BUTTON = QPushButton("Save")
        SELF.SAVE_BUTTON.setStyleSheet("""
            QPushButton {
                background-color: #39FF14;  /* Neon green background */
                border: 2px solid #39FF14;  /* Neon green border */
                color: #000000;  /* Black text */
                padding: 10px;
                border-radius: 10px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #FF0000;  /* Red background on hover */
                border: 2px solid #FF0000;  /* Red border on hover */
                color: #FFFFFF;  /* White text on hover */
            }
        """)
        SELF.SAVE_BUTTON.clicked.connect(SELF.save_reports)
        BUTTON_LAYOUT.addWidget(SELF.SAVE_BUTTON)

        # Open Button
        SELF.OPEN_BUTTON = QPushButton("Open")
        SELF.OPEN_BUTTON.setStyleSheet("""
            QPushButton {
                background-color: #39FF14;  /* Neon green background */
                border: 2px solid #39FF14;  /* Neon green border */
                color: #000000;  /* Black text */
                padding: 10px;
                border-radius: 10px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #FF0000;  /* Red background on hover */
                border: 2px solid #FF0000;  /* Red border on hover */
                color: #FFFFFF;  /* White text on hover */
            }
        """)
        SELF.OPEN_BUTTON.clicked.connect(SELF.open_reports)
        BUTTON_LAYOUT.addWidget(SELF.OPEN_BUTTON)

        # Delete Button
        SELF.DELETE_BUTTON = QPushButton("Delete")
        SELF.DELETE_BUTTON.setStyleSheet("""
            QPushButton {
                background-color: #FF0000;  /* Red background */
                border: 2px solid #FF0000;  /* Red border */
                color: #FFFFFF;  /* White text */
                padding: 10px;
                border-radius: 10px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: blue;  /* Neon green background on hover */
                border: 2px solid blue;  /* Neon green border on hover */
                color: #000000;  /* Black text on hover */
            }
        """)
        SELF.DELETE_BUTTON.clicked.connect(SELF.delete_bug)
        SELF.DELETE_BUTTON.setEnabled(False)
        BUTTON_LAYOUT.addWidget(SELF.DELETE_BUTTON)

        FORM_LAYOUT.addLayout(BUTTON_LAYOUT)
        SELF.form_widget.setLayout(FORM_LAYOUT)
        SELF.splitter.addWidget(SELF.form_widget)


        SELF.REPORT_TABLE = QTableWidget()
        SELF.REPORT_TABLE.setColumnCount(6)
        SELF.REPORT_TABLE.setHorizontalHeaderLabels(["ID", "Name", "Title", "Severity", "Status", "Date"])
        SELF.REPORT_TABLE.setColumnHidden(0, True)
        SELF.REPORT_TABLE.horizontalHeader().setStretchLastSection(True)
        SELF.REPORT_TABLE.setSelectionBehavior(QTableWidget.SelectRows)
        SELF.REPORT_TABLE.cellClicked.connect(SELF.select_bug)

        # Disable column header editing
        SELF.REPORT_TABLE.setEditTriggers(QTableWidget.NoEditTriggers)

        # Apply a cooler, more modern style with neon green and black
        SELF.REPORT_TABLE.setStyleSheet("""
        QTableWidget {
            background-color: #2D2D2D;  /* Dark background for a sleek look */
            border: 1px solid #444444;  /* Dark border for a subtle contrast */
            color: #F0F0F0;  /* Light text color */
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;  /* Modern font */
            font-size: 12px;
            selection-background-color: #39FF14;  /* Neon green selection for better contrast */
            selection-color: #000000;  /* Black text on selected rows */
        }

        QTableWidget::item {
            border-bottom: 1px solid #444444;  /* Light border between rows */
        }

        QTableWidget::item:selected {
            background-color: #39FF14;  /* Neon green background for selected rows */
            border: 1px solid #00FF00;  /* Highlight the border with a brighter neon green */
        }

        QHeaderView::section {
            background-color: #333333;  /* Darker header background */
            color: #FFFFFF;  /* White text for the header */
            padding: 5px;
            font-size: 13px;
            font-weight: bold;
        }

        QHeaderView::section:horizontal {
            border: none;
        }

        QTableWidget::indicator {
            background-color: #444444;  /* Dark background for checkboxes */
        }

        QTableWidget::indicator:checked {
            background-color: #39FF14;  /* Neon green color for checked indicator */
        }
        """)

        SELF.splitter.addWidget(SELF.REPORT_TABLE)

        layout.addWidget(SELF.splitter)

        SELF.STATUS_BAR = QStatusBar()
        SELF.setStatusBar(SELF.STATUS_BAR)

        SELF.central_widget.setLayout(layout)


    def init_database(SELF):
        SELF.conn = sqlite3.connect("bug_reports.db")
        SELF.cursor = SELF.conn.cursor()

        try:
            SELF.cursor.execute("""
                ALTER TABLE bugs ADD COLUMN name TEXT;
            """)
        except sqlite3.OperationalError:
            pass

        SELF.cursor.execute("""
            CREATE TABLE IF NOT EXISTS bugs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT,
                title TEXT,
                description TEXT,
                severity TEXT,
                status TEXT DEFAULT 'Open',
                submission_date TEXT
            )
        """)
        SELF.conn.commit()
        SELF.load_reports()

    def submit_bug(SELF):
        name = SELF.NAME_INPUT.text().strip()
        title = SELF.TITLE_INPUT.text().strip()
        description = SELF.DESCRIPTION_INPUT.toPlainText().strip()
        severity = SELF.SEVERITY_DROPDOWN.currentText()
        submission_date = SELF.DATE_INPUT.date().toString("yyyy-MM-dd")

        if not name or not title or not description:
            SELF.STATUS_BAR.showMessage("Error: Name, Title, and Description are required.", 3000)
            return

        if SELF.selected_bug_id:
            SELF.cursor.execute("""
                UPDATE bugs SET name = ?, title = ?, description = ?, severity = ?, submission_date = ?
                WHERE id = ?
            """, (name, title, description, severity, submission_date, SELF.selected_bug_id))
            SELF.selected_bug_id = None
            SELF.STATUS_BAR.showMessage("Bug updated successfully.", 3000)
        else:
            SELF.cursor.execute("""
                INSERT INTO bugs (name, title, description, severity, submission_date)
                VALUES (?, ?, ?, ?, ?)
            """, (name, title, description, severity, submission_date))
            SELF.STATUS_BAR.showMessage("Bug submitted successfully.", 3000)

        SELF.conn.commit()
        SELF.load_reports()

        SELF.NAME_INPUT.clear()
        SELF.TITLE_INPUT.clear()
        SELF.DESCRIPTION_INPUT.clear()
        SELF.SEVERITY_DROPDOWN.setCurrentIndex(0)
        SELF.DATE_INPUT.setDate(QDate.currentDate())
        SELF.MARK_DONE_BUTTON.setEnabled(False)

    def load_reports(SELF):
        SELF.REPORT_TABLE.setRowCount(0)
        SELF.cursor.execute("SELECT id, name, title, severity, status, submission_date FROM bugs")
        for row_index, row_data in enumerate(SELF.cursor.fetchall()):
            SELF.REPORT_TABLE.insertRow(row_index)
            for col_index, col_data in enumerate(row_data):
                SELF.REPORT_TABLE.setItem(row_index, col_index, QTableWidgetItem(str(col_data)))

    def select_bug(SELF, row, column):
        # Check if the same row is clicked again, if so, deselect it
        bug_id = int(SELF.REPORT_TABLE.item(row, 0).text())
        
        if SELF.selected_bug_id == bug_id:  # Deselect if it's the same bug
            SELF.selected_bug_id = None
            SELF.clear_inputs()
            SELF.MARK_DONE_BUTTON.setEnabled(False)
            SELF.DELETE_BUTTON.setEnabled(False)
            SELF.REPORT_TABLE.clearSelection()  # Deselect the row
        else:
            SELF.selected_bug_id = bug_id
            SELF.cursor.execute("SELECT name, title, description, severity, submission_date FROM bugs WHERE id = ?", 
                                (SELF.selected_bug_id,))
            bug = SELF.cursor.fetchone()
            SELF.NAME_INPUT.setText(bug[0])
            SELF.TITLE_INPUT.setText(bug[1])
            SELF.DESCRIPTION_INPUT.setText(bug[2])
            SELF.SEVERITY_DROPDOWN.setCurrentText(bug[3])
            SELF.DATE_INPUT.setDate(QDate.fromString(bug[4], "yyyy-MM-dd"))
            SELF.MARK_DONE_BUTTON.setEnabled(True)
            SELF.DELETE_BUTTON.setEnabled(True)
    def clear_inputs(SELF):
        # Helper method to clear input fields
        SELF.NAME_INPUT.clear()
        SELF.TITLE_INPUT.clear()
        SELF.DESCRIPTION_INPUT.clear()
        SELF.SEVERITY_DROPDOWN.setCurrentIndex(0)
        SELF.DATE_INPUT.setDate(QDate.currentDate())
    def mark_bug_done(SELF):
        if SELF.selected_bug_id:
            SELF.cursor.execute("UPDATE bugs SET status = 'Done' WHERE id = ?", (SELF.selected_bug_id,))
            SELF.conn.commit()
            SELF.load_reports()
            SELF.STATUS_BAR.showMessage("Bug marked as done.", 3000)
            SELF.MARK_DONE_BUTTON.setEnabled(False)
            SELF.selected_bug_id = None
            SELF.DELETE_BUTTON.setEnabled(False)

        SELF.NAME_INPUT.clear()
        SELF.TITLE_INPUT.clear()
        SELF.DESCRIPTION_INPUT.clear()
        SELF.SEVERITY_DROPDOWN.setCurrentIndex(0)
        SELF.DATE_INPUT.setDate(QDate.currentDate())

    def delete_bug(SELF):
        if SELF.selected_bug_id:
            reply = SELF.show_confirmation("Are you sure you want to delete this bug report?")
            if reply:
                SELF.cursor.execute("DELETE FROM bugs WHERE id = ?", (SELF.selected_bug_id,))
                SELF.conn.commit()
                SELF.load_reports()
                SELF.STATUS_BAR.showMessage("Bug report deleted.", 3000)
                SELF.selected_bug_id = None
                SELF.DELETE_BUTTON.setEnabled(False)

                SELF.NAME_INPUT.clear()
                SELF.TITLE_INPUT.clear()
                SELF.DESCRIPTION_INPUT.clear()
                SELF.SEVERITY_DROPDOWN.setCurrentIndex(0)
                SELF.DATE_INPUT.setDate(QDate.currentDate())

    def show_confirmation(SELF, message):
        from PyQt5.QtWidgets import QMessageBox
        MSG_BOX = QMessageBox(SELF)
        MSG_BOX.setIcon(QMessageBox.Question)
        MSG_BOX.setText(message)
        MSG_BOX.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        reply = MSG_BOX.exec_()
        return reply == QMessageBox.Yes

    def save_reports(SELF):
        # Create a file dialog to select save location and file format
        file_format, _ = QFileDialog.getSaveFileName(
            SELF,
            "Save Reports",
            "",
            "Text Files (*.txt);;Word Documents (*.docx);;JSON Files (*.json);;All Files (*)"
        )
        
        if file_format:
            # Check the file format and save accordingly
            if file_format.endswith(".txt"):
                SELF.save_as_text(file_format)
            elif file_format.endswith(".docx"):
                SELF.save_as_docx(file_format)
            elif file_format.endswith(".json"):
                SELF.save_as_json(file_format)

            SELF.STATUS_BAR.showMessage(f"Reports saved to {file_format}", 3000)

    def save_as_text(SELF, file_path):
        """Saves the bug reports as a plain text file."""
        with open(file_path, "w") as f:
            SELF.cursor.execute("SELECT name, title, description, severity, status, submission_date FROM bugs")
            for bug in SELF.cursor.fetchall():
                f.write(f"Name: {bug[0]}\n")
                f.write(f"Title: {bug[1]}\n")
                f.write(f"Description: {bug[2]}\n")
                f.write(f"Severity: {bug[3]}\n")
                f.write(f"Status: {bug[4]}\n")
                f.write(f"Submission Date: {bug[5]}\n")
                f.write("-" * 50 + "\n")

    def save_as_docx(self, file_path):
        doc = Document()
        
        title = doc.add_paragraph()
        run = title.add_run("BUG REPORTS")
        run.font.size = 240000  
        run.font.color.rgb = RGBColor(57, 255, 20)  
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT  

        doc.add_paragraph("-" * 50)
        
        self.cursor.execute("SELECT name, title, description, severity, status, submission_date FROM bugs")
        for bug in self.cursor.fetchall():
            doc.add_paragraph(f"Name: {bug[0]}")
            doc.add_paragraph(f"Title: {bug[1]}")
            doc.add_paragraph(f"Description: {bug[2]}")
            doc.add_paragraph(f"Severity: {bug[3]}")
            doc.add_paragraph(f"Status: {bug[4]}")
            doc.add_paragraph(f"Submission Date: {bug[5]}")
            doc.add_paragraph("-" * 50)
        
        doc.save(file_path)

    def save_as_json(SELF, file_path):
        """Saves the bug reports as a JSON file."""
        import json
        reports = []
        SELF.cursor.execute("SELECT name, title, description, severity, status, submission_date FROM bugs")
        for bug in SELF.cursor.fetchall():
            reports.append({
                "name": bug[0],
                "title": bug[1],
                "description": bug[2],
                "severity": bug[3],
                "status": bug[4],
                "submission_date": bug[5]
            })
        
        with open(file_path, "w") as f:
            json.dump(reports, f, indent=4)

    def open_reports(SELF):
        file_path, _ = QFileDialog.getOpenFileName(SELF, "Open Reports", "", "Text Files (*.txt);;All Files (*)")
        if file_path:
            with open(file_path, "r") as f:
                content = f.read()
                SELF.STATUS_BAR.showMessage(f"Opened file {file_path}", 3000)
                print(content)  
if __name__ == "__main__":
    app = QApplication(sys.argv)
    mainWin = BugReportApp()
    mainWin.show()
    sys.exit(app.exec())